from dataclasses import dataclass, field
from typing import List, Any
import copy
import json

from tools.llm.api_client_qwen_openai import *
import openai
openai.api_base = "http://116.62.63.204:8001/v1"

import fitz
from fitz import TextPage

import docx
from docx import Document   # api.Document
from docx.oxml.ns import qn


import docx  # 导入python-docx库
# from docx.document import Document  # document.Document
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.table import _Cell, Table
from docx.text.paragraph import Paragraph

# llm = LLM_Qwen()

from enum import Enum, auto

LLM_Doc_DEBUG = False
def dprint(*args, **kwargs):
    if LLM_Doc_DEBUG:
        print(*args, **kwargs)

def is_win():
    import platform
    sys_platform = platform.platform().lower()
    if "windows" in sys_platform:
        return True
    else:
        return False

# =========================================管理doc的层次化递归数据结构=================================================
from utils.Hierarchy_Node import Hierarchy_Node
@dataclass
class Image_Data():
    name: str
    data: str
    width: int
    height: int

@dataclass
class Table_Data():
    # table标题
    index:      str = ''  # '2.1.6.3.1-1'
    head:       str = ''  # '南麂微电网2015年各支路潮流及各配变负载率计算结果'
    unit:       str = ''  # '万元'
    annotate:   str = ''  # '注: ...'
    # table内容
    text:       str = ''  # 表格内容文本
    obj:        Any = None  # 表格的Table对象

@dataclass
class Doc_Node_Data():
    type:str = ''   # 'text' 'table' 'image'
    text:str = ''
    table:Table_Data = None    # Table_Data 对象
    image:Image_Data = None

@dataclass
class Doc_Node_Content():
    level: int
    name: str       # 如: '1.1.3'
    heading: str    # 如: '建设必要性'
    data_list:List[Doc_Node_Data] = field(default_factory=list)   # 元素包括(text:str, image:Image_Part, table:str)
    # text: str       # 如: '本项目建设是必要的...'
    # image: Image_Part
# =========================================管理doc的层次化递归数据结构=================================================

# LLM_Doc：采用python-docx解析文档，采用win32com解决页码问题
class LLM_Doc():
    def __init__(self, in_file_name, in_llm=None):
        self.doc_name = in_file_name    # 文件名

        self.win32_doc = None
        self.win32_doc_app = None
        self.win32_constant = None

        self.doc = None                 # Document对象
        self.doc_root = None            # 存放doc层次化数据
        self.doc_root_parsed = False    # 是否已经解析为层次化数据

        self.llm = in_llm
        if self.llm is None:
            self.llm = LLM_Qwen(
                history=False,
                # history_max_turns=50,
                # history_clear_method='pop',
                temperature=0.7,
                url='http://127.0.0.1:8001/v1',
                need_print=False,
            )

        try:
            self.doc = Document(self.doc_name)
        except Exception as e:
            print(f'文件"{self.doc_name}" 未找到。')

    def ask_docx(self, in_query, in_max_level=3):
        file = self.doc_name
        doc = self
        doc.parse_all_docx()
        toc = doc.get_toc_list_json_string(in_max_level=in_max_level)

        # -------------------------------找到query内容所在章节---------------------------------------
        prompt = '''
        这是文档的目录结构"{toc}",
        请问这个问题"{query}"涉及的内容应该在具体的哪个章节中，不解释，请直接以"章节编号"形式返回。
        '''
        prompt = prompt.format(toc=toc, query=in_query)
        self.llm.need_print = False         # 关闭print输出
        res = self.llm.ask_prepare(prompt).get_answer_and_sync_print()

        # --------------将'1.1.3 some章节'转换为'1.1.3'----------------------
        re_result = re.search(r"\d+(.\d+)*", res).group(0)

        # --------------获取'1.1.3'对应章节下的text_got----------------------
        node = doc.find_doc_root(re_result)
        inout_text = []
        doc.get_text_from_doc_node(inout_text, node)
        text_got = '\n'.join(inout_text)

        # --------------对text_got进行限定范围的query----------------------
        prompt2 = '''
        请根据材料"{text_got}"中的内容, 回答问题"{query}"。
        '''

        prompt2 = prompt2.format(text_got=text_got, query=in_query)
        print(prompt2)
        print(f'材料长度为: {len(text_got)}')
        # self.llm.need_print = True          # 打开print输出
        gen = self.llm.ask_prepare(prompt2).get_answer_generator()
        return gen

    def win32com_init(self):
        print(f'win32尝试打开文件"{self.doc_name}"')
        if is_win():
            import win32com.client as win32
            from win32com.client import constants
            self.win32_constant = constants
        else:
            print(f"未在windows系统下运行，无法获取word文档页码.")
            return

        import os

        self.win32_doc_app = win32.gencache.EnsureDispatch('Word.Application')  # 打开word应用程序
        self.win32_doc_app.Visible = False
        # doc_app.Visible = True
        curr_path = os.getcwd()
        self.win32_doc = self.win32_doc_app.Documents.Open(self.doc_name)
        print(f'win32打开了文件"{self.doc_name}"')

    def win32_close_file(self):
        if self.win32_doc:
            try:
                self.win32_doc_app.Documents.Close(self.doc_name)
            except Exception as e:
                print(f'关闭文件"{self.doc_name}"出错: {e}')

    def iter_block_items(self, parent):
        """
        Yield each paragraph and table child within *parent*, in document order.
        Each returned value is an instance of either Table or Paragraph. *parent*
        would most commonly be a reference to a main Document object, but
        also works for a _Cell object, which itself can contain paragraphs and tables.
        """
        if isinstance(parent, docx.document.Document):
            parent_elm = parent.element.body
        elif isinstance(parent, _Cell):
            parent_elm = parent._tc
        else:
            raise ValueError("something's not right")

        for child in parent_elm.iterchildren():
            if isinstance(child, CT_P):
                yield Paragraph(child, parent)
            elif isinstance(child, CT_Tbl):
                yield Table(child, parent)

    # 获取node下目录(table of content)的json格式, list形式，节省字符串长度
    def get_toc_list_json_string(self, in_max_level=3):
        import json
        toc = []
        if self.doc_root is None:
            return json.dumps([], indent=2, ensure_ascii=False)

        self.doc_root.get_toc_list_json(toc, self.doc_root, in_max_level)

        return json.dumps(toc, indent=2, ensure_ascii=False)

    # 获取node下目录(table of content)的json格式, dict形式，比较占用字符串长度
    def get_toc_dict_json_string(self, in_max_level=3):
        import json
        toc = {}
        if self.doc_root is None:
            return json.dumps({}, indent=2, ensure_ascii=False)

        self.doc_root.get_toc_dict_json(toc, self.doc_root, in_max_level)

        return json.dumps(toc, indent=2, ensure_ascii=False)

    # 用'1.1.3'这样的字符串查找node
    def find_doc_root(self, in_node_s):
        if self.doc_root is None:
            return None

        return self.doc_root.find(in_node_s)

    # 遍历打印node下的所有内容
    def print_from_doc_node(self, in_node='root'):   # in_node为'1.1.3'或者Hierarchy_Node对象
        # 如果输入'1.1.3'这样的字符串
        if type(in_node)==str:
            node_s = in_node
            in_node = self.doc_root.find(node_s)
            if in_node is None:
                print(f'节点"{node_s}"未找到.')
                return

        # 如果输入Hierarchy_Node对象
        if in_node is None:
            return
        else:
            node = in_node

        node_level = node.node_data.level
        node_name = node.node_data.name
        node_heading = node.node_data.heading

        # 获取node文本
        node_content = self.get_node_data_callback(node)

        print(f'{"-"*node_level}{node_name}-{node_heading}{"-"*(80-node_level-len(node_name)-len(node_heading))}')
        print(f'{node_content}')

        for child in node.children:
            self.print_from_doc_node(child)

    # 遍历输出node下的所有内容
    def get_text_from_doc_node(self, inout_text_list, in_node='root'):   # in_node为'1.1.3'或者Hierarchy_Node对象
        # 如果输入'1.1.3'这样的字符串
        if type(in_node)==str:
            node_s = in_node
            in_node = self.doc_root.find(node_s)
            if in_node is None:
                dprint(f'节点"{node_s}"未找到.')
                return

        # 如果输入Hierarchy_Node对象
        if in_node is None:
            return
        else:
            node = in_node

        node_level = node.node_data.level
        node_name = node.node_data.name
        node_heading = node.node_data.heading

        # 获取node文本
        node_content = self.get_node_data_callback(node)

        inout_text_list.append(f'{"-"*node_level}{node_name}-{node_heading}{"-"*(80-node_level-len(node_name)-len(node_heading))}')
        inout_text_list.append(f'{node_content}')

        for child in node.children:
            self.get_text_from_doc_node(inout_text_list, child)

    # 为table增加annotate'注: ...'
    def get_table_annotate_from_text(self, in_text, inout_table_obj):
        match_annotate = re.search(r'\s*(?<=注[:：]).*', in_text)
        match_annotate = match_annotate.group(0).strip() if match_annotate else ''
        inout_table_obj.annotate = match_annotate
        return inout_table_obj

    # 从文本中解析出table的标题
    def get_table_head_from_text(self, in_text, inout_table_obj):
        '\n附表2.1.6.3.1 -1   南麂岛10kV线路总概算表   单位：万元'
        match_index = re.search(r'^((\s*)(表|附表))[\d\.\s-]*', in_text)  # 开头可能有'\n'这种字符，2.1.6.3.1和-1中间可能有空格' '
        match_index = match_index.group(0).strip() if match_index else ''

        match_unit = re.search(r'(?<=单位[：|:])\s*[\u4e00-\u9fa5]+', in_text)     # (?<=exp)匹配exp后面，(?=exp)匹配exp前面，(?!exp)匹配后面跟的不是exp，(?<!exp)匹配前面不是exp
        match_unit = match_unit.group(0).strip() if match_unit else ''

        match_unit_all = re.search(r'单位[：|:]\s*[\u4e00-\u9fa5]+', in_text)     # (?<=exp)匹配exp后面，(?=exp)匹配exp前面，(?!exp)匹配后面跟的不是exp，(?<!exp)匹配前面不是exp
        match_unit_all = match_unit_all.group(0).strip() if match_unit_all else ''

        match_head = in_text.replace(match_index, '').replace(match_unit_all, '').strip()

        inout_table_obj.index = match_index     # '附表21'
        inout_table_obj.head = match_head       # '南麂岛10kV线路总概算表'
        inout_table_obj.unit = match_unit       # '万元'
        # inout_table_obj.annotate= ''        # '注: ...'

        return inout_table_obj

    # 替代para.text，因为para.text解析word xml的文字时，会漏掉'w:smartTag'或'w:ins'下的'w:r'，导致如'表2.1.6.3.4'变为'表.3.4'的问题
    def _patch_get_text(self, in_para):
        text = []
        for elem in in_para._element:
            if elem.tag == qn('w:r'):
                text.append(elem.text)
            elif elem.tag == qn('w:ins') or elem.tag == qn('w:smartTag'):
                for sub_elem in elem:
                    if sub_elem.tag == qn('w:r'):
                        text.append(sub_elem.text)
        return ''.join(text)

    # 将doc解析为层次结构，每一级标题（容器）下都有text、table、image等对象
    def parse_all_docx(self):
        # 获取node上一级即parent的name，如'1.1.3'的上一级为'1.1', '1'的上一级为'root'
        def find_parent_node(in_node_name):
            dprint(f'================查找节点"{in_node_name}"的父节点', end='')
            if len(in_node_name.split('.')) == 1:
                parent_name = 'root'
                dprint(f'"{parent_name}"========')
                return self.doc_root.find(parent_name)
            else:
                # name_list = in_node_name.split('.')
                # name_list.pop()
                # parent_name = '.'.join(name_list)
                # print(f'"{parent_name}"========')
                # return self.doc_root.find(parent_name)

                name_list = in_node_name.split('.')
                while True:
                    # 循环pop()，直到找到parent_node，例如3.4后面突然出现3.4.1.1，这时候的parent_node就3.4而不是3.4.1
                    name_list.pop()
                    parent_name = '.'.join(name_list)
                    dprint(f'"{parent_name}"========')
                    node = self.doc_root.find(parent_name)
                    if node:
                        return node

        # 处理root
        self.doc_root = Hierarchy_Node(Doc_Node_Content(0, 'root', 'root根'))

        current_node = self.doc_root
        current_node_name = 'root'
        current_level = 0

        # 上一个文本，主要用于查找table标题
        last_para_text = ''
        # 用于判断'注：...'前面是否为表格
        last_is_table = False
        last_table_obj = None

        # 递归遍历doc
        # for para in self.doc.paragraphs:
        for block in self.iter_block_items(self.doc):
            if block.style.name == 'Normal Table':
                #--------------------------------------------找到表格-----------------------------------------------------
                table = block
                inout_table_obj = Table_Data(obj=table)
                self.get_table_head_from_text(last_para_text, inout_table_obj)  # 解析table的标题
                # table_obj = Table_Data(head=self.get_table_head_from_text(last_para_text), obj=table)
                # 添加内容(text、table、image等元素)
                self.parse_node_data_callback(
                    in_node_data_list_ref=current_node.node_data.data_list,
                    in_data=Doc_Node_Data(type='table', table=inout_table_obj)
                )  # 这里para.text是文本内容 Doc_Node_Data.data_list中的text

                last_is_table = True
                last_table_obj = inout_table_obj
            else:
                #--------------------------------------------找到标题head、文本text、图image等------------------------------
                para = block
                if last_is_table:
                    self.get_table_annotate_from_text(self._patch_get_text(para), last_table_obj)
                    # *******************测试para.text的bug************************
                    # if '注：由于各配电变压器' in para.text:
                    #     print(f'注：由于各配电变压器: [{self._patch_get_text(para)}]')
                    #     print(f'注：由于各配电变压器: [{para.text}]')
                        # print(f'注：由于各配电变压器: [{para._element.xml}]')
                        # print(f'annotate: {last_table_obj.annotate}')
                    # *******************测试para.text的bug************************

                last_para_text = self._patch_get_text(para)
                style = para.style.name                     # "Heading 1"
                style_name  = style.split(' ')[0]           # 标题 "Heading

                if style_name=='Heading':
                    # ----------------------------------找到标题heading，处理Hierarchy层次结构-------------------------------
                    new_level = int(style.split(' ')[1])    # 标题级别 1

                    # 计算current_node_name, 如：'1.1.1'或'1.2'
                    if current_node_name=='root':
                        current_node_name = '.'.join(['1']*new_level)   # 1级为'1', 如果直接为2级就是'1.1'
                    else:
                        if new_level == current_level:
                            dprint(f'----------------------------current_node_name: {current_node_name}----------------------------')
                            dprint(f'new_level: {new_level}')
                            dprint(f'current_level: {current_level}')
                            # ‘1.1.1’变为'1.1.2'
                            new_node_list = current_node_name.split('.')  # ['1', '1', '1']
                            last_num = int(new_node_list[-1]) + 1  # 2
                            new_node_list.pop()  # ['1', '1']
                            current_node_name = '.'.join(new_node_list) + '.' + str(last_num)  # '1.1.2'
                            # current_node_name = current_node_name[:-1] + str(int(current_node_name[-1])+1)
                        elif new_level > current_level:
                            dprint(f'----------------------------current_node_name: {current_node_name}----------------------------')
                            dprint(f'new_level: {new_level}')
                            dprint(f'current_level: {current_level}')
                            # ‘1.1.1’变为'1.1.1.1.1'
                            current_node_name += '.' + '.'.join(['1']*(new_level-current_level))
                        elif new_level < current_level:
                            dprint(f'----------------------------current_node_name: {current_node_name}----------------------------')
                            dprint(f'new_level: {new_level}')
                            dprint(f'current_level: {current_level}')
                            # ‘1.1.1’变为'1.2' 或 ‘1.1.1’变为'2'
                            new_node_list = current_node_name.split('.')    # ['1', '1', '1']
                            for i in range(current_level-new_level):
                                new_node_list.pop()                         # ['1', '1']
                            last_num = int(new_node_list[-1]) +1                      # 2
                            new_node_list.pop()                             # ['1']
                            if len(new_node_list)>0:
                                current_node_name = '.'.join(new_node_list) + '.' + str(last_num)   # '1.2'
                            else:
                                current_node_name = str(last_num)
                            # current_node_name = current_node_name[:-1-2*(current_level-new_level)] + str(int(current_node_name[-1-2*(current_level-new_level)])+1)
                    current_level = new_level

                    # 找到parent节点，并添加new_node
                    new_node = Hierarchy_Node(Doc_Node_Content(current_level, current_node_name, self._patch_get_text(para))) # 这里para.text是标题内容 Doc_Node_Data.heading
                    parent_node = find_parent_node(current_node_name)
                    parent_node.add_child(new_node)

                    # 刷新current状态
                    current_node = new_node
                else:
                    # ------------------------------------------找到内容：text、image等------------------------------------
                    self.parse_node_data_callback(
                        in_node_data_list_ref=current_node.node_data.data_list,
                        in_data=Doc_Node_Data(type='text', text=self._patch_get_text(para))
                    )    # 这里para.text是文本内容 Doc_Node_Data.data_list中的text

                last_is_table = False
                last_table_obj = None

        self.doc_root_parsed = True

    # 解析node数据的callback
    def parse_node_data_callback(self, in_node_data_list_ref, in_data):
        in_node_data_list_ref.append(in_data)

    # 读取node数据的callback
    def get_node_data_callback(self, in_node):
        node_content = ''

        for node_data in in_node.node_data.data_list:
            if node_data.type=='text':
                # 普通文本
                node_content += node_data.text + '\n'
            elif node_data.type=='table':
                # 表格
                # table头
                tbl_index = node_data.table.index
                tbl_head = node_data.table.head
                tbl_unit = node_data.table.unit
                tbl_annotate = node_data.table.annotate
                half1 = (80-len(tbl_index)-len(tbl_head)-len(tbl_unit)-len('表格[::]'))//2
                half2 = (80-len(tbl_annotate)-len('[注: ]'))//2

                node_content += '-'*half1 + f'表格[{tbl_index}:{tbl_head}:{tbl_unit}]' + '-'*half1 + '\n'

                # table内容
                for row in node_data.table.obj.rows:
                    for cell in row.cells:
                        node_content += cell.text + '\t'
                    node_content += '\n'
                # table注解
                if tbl_annotate:
                    node_content += '-'*half2 + f'[注: {tbl_annotate}]' + '-'*half2 + '\n'
                else:
                    node_content += '-'*80
            else:
                pass

        return node_content

    # 将pdf解析为层次结构，每一级标题（容器）下都有text、image等对象
    def parse_all_pdf(self):
        # 获取上一级的name，如'1.1.3'的上一级为'1.1', '1'的上一级为'root'
        def find_parent_node(in_node_name):
            dprint(f'================查找节点"{in_node_name}"的父节点', end='')
            if len(in_node_name.split('.')) == 1:
                parent_name = 'root'
                dprint(f'"{parent_name}"========')
                return self.doc_root.find(parent_name)
            else:
                # name_list = in_node_name.split('.')
                # name_list.pop()
                # parent_name = '.'.join(name_list)
                # print(f'"{parent_name}"========')
                # return self.doc_root.find(parent_name)

                name_list = in_node_name.split('.')
                while True:
                    # 循环pop()，直到找到parent_node，例如3.4后面突然出现3.4.1.1，这时候的parent_node就3.4而不是3.4.1
                    name_list.pop()
                    parent_name = '.'.join(name_list)
                    dprint(f'"{parent_name}"========')
                    node = self.doc_root.find(parent_name)
                    if node:
                        return node

        # 处理root
        self.doc_root = Hierarchy_Node(Doc_Node_Content(0, 'root', 'root根'))

        current_node = self.doc_root
        current_node_name = 'root'
        current_level = 0

        # 递归遍历doc
        for para in self.doc.paragraphs:
            style = para.style.name                     # "Heading 1"
            style_name  = style.split(' ')[0]           # 标题 "Heading

            if style_name=='Heading':
                # 标题
                new_level = int(style.split(' ')[1])    # 标题级别 1

                # 计算current_node_name, 如：'1.1.1'或'1.2'
                if current_node_name=='root':
                    current_node_name = '.'.join(['1']*new_level)   # 1级为'1', 如果直接为2级就是'1.1'
                else:
                    if new_level == current_level:
                        dprint(f'----------------------------current_node_name: {current_node_name}----------------------------')
                        dprint(f'new_level: {new_level}')
                        dprint(f'current_level: {current_level}')
                        # ‘1.1.1’变为'1.1.2'
                        new_node_list = current_node_name.split('.')  # ['1', '1', '1']
                        last_num = int(new_node_list[-1]) + 1  # 2
                        new_node_list.pop()  # ['1', '1']
                        current_node_name = '.'.join(new_node_list) + '.' + str(last_num)  # '1.1.2'
                        # current_node_name = current_node_name[:-1] + str(int(current_node_name[-1])+1)
                    elif new_level > current_level:
                        dprint(f'----------------------------current_node_name: {current_node_name}----------------------------')
                        dprint(f'new_level: {new_level}')
                        dprint(f'current_level: {current_level}')
                        # ‘1.1.1’变为'1.1.1.1.1'
                        current_node_name += '.' + '.'.join(['1']*(new_level-current_level))
                    elif new_level < current_level:
                        dprint(f'----------------------------current_node_name: {current_node_name}----------------------------')
                        dprint(f'new_level: {new_level}')
                        dprint(f'current_level: {current_level}')
                        # ‘1.1.1’变为'1.2' 或 ‘1.1.1’变为'2'
                        new_node_list = current_node_name.split('.')    # ['1', '1', '1']
                        for i in range(current_level-new_level):
                            new_node_list.pop()                         # ['1', '1']
                        last_num = int(new_node_list[-1]) +1                      # 2
                        new_node_list.pop()                             # ['1']
                        if len(new_node_list)>0:
                            current_node_name = '.'.join(new_node_list) + '.' + str(last_num)   # '1.2'
                        else:
                            current_node_name = str(last_num)
                        # current_node_name = current_node_name[:-1-2*(current_level-new_level)] + str(int(current_node_name[-1-2*(current_level-new_level)])+1)
                current_level = new_level

                # 找到parent节点，并添加new_node
                new_node = Hierarchy_Node(Doc_Node_Content(current_level, current_node_name, self._patch_get_text(para))) # 这里para.text是标题内容 Doc_Node_Data.heading
                parent_node = find_parent_node(current_node_name)
                parent_node.add_child(new_node)

                # 刷新current状态
                current_node = new_node
            else:
                # --------------------------------------------内容：text、image等-------------------------------------
                self.parse_node_data_callback(
                    in_node_data_list_ref=current_node.node_data.data_list,
                    in_data=Doc_Node_Data(type='text', text=self._patch_get_text(para))
                )  # 这里para.text是文本内容 Doc_Node_Data.data_list中的text

        self.doc_root_parsed = True

    def get_para_inline_images(self, in_para):
        image = {
            'name':'',
            'data':None,
            'width':0,
            'height':0,
        }
        images_list = []

        # 打印ImagePart的成员函数
        # for item in dir(docx.parts.image.ImagePart):
        #     print(item)

        # 这一段由python-docx库的issue中网友mfripp提供：https://github.com/python-openxml/python-docx/issues/249
        for run in in_para.runs:
            for inline in run._r.xpath("w:drawing/wp:inline"):
                width = float(inline.extent.cx)  # in EMUs https://startbigthinksmall.wordpress.com/2010/01/04/points-inches-and-emus-measuring-units-in-office-open-xml/
                height = float(inline.extent.cy)
                if inline.graphic.graphicData.pic is not None:
                    rId = inline.graphic.graphicData.pic.blipFill.blip.embed
                    image_part = self.doc.part.related_parts[rId]
                    filename = image_part.filename      # 文件名称(其实是类型), 如"image.wmf"
                    bytes_of_image = image_part.blob    # 文件数据(bytes)
                    image['name'] = filename
                    image['data'] = bytes_of_image
                    image['width'] = width
                    image['height'] = height
                    images_list.append(copy.deepcopy(image))

                    # print(f'image_part:{image_part}, width:{width}, height:{height}, rId: {rId}, image:{image}, filename:{filename}')
                    # with open('a.wmf', 'wb') as f:  # make a copy in the local dir
                    #     f.write(bytes_of_image)
        return images_list


    def get_all_inline_images(self):
        images_list = []

        for para in self.doc.paragraphs:
            images_list += self.get_para_inline_images(para)
        return images_list

    def get_paras(self):
        for para in self.doc.paragraphs:
            yield para.text

    # 检查某doc文档所有para的错别词
    def check_wrong_written_in_docx_file(self, in_llm):
        try:
            doc = Document(self.doc_name)
        except Exception as e:
            print(f'文件"{self.doc_name}" 未找到。')
            return

        ii = 0
        result_list = []
        for para in self.get_paras():
            ii += 1
            # print(f"正在分析第{ii}个段落...")

            # para result登记
            print("*"*30+f"正在分析第{ii}个段落"+"*"*30)
            check_prompt = f'现在帮我在这段文字中找错别字："{para}"'
            print(f'本段落内容："{para}"')
            # in_llm.print_history()
            # print(f'check prompt 为：{check_prompt}')
            result_list.append(in_llm.ask_prepare(check_prompt).get_answer_and_sync_print())

        return result_list

# ============================关于角色提示============================
# 一、你希望llm了解哪些信息：
# 1) where are you based?
# 2) what do you do for work?
# 3) what are your hobbies and interests?
# 4) what subject can you talk about for hours?
# 5) what are some goals you have?
# 二、你希望llm怎样回复：
# 1) how formal or casual should llm be?
# 2) how long or short should responses generally be?
# 3) how do you want to be addressed?
# 4) should llm have opinions on topics or remain neutral?
# ============================关于角色提示============================
def main1():
    llm = LLM_Qwen(
        url='http://116.62.63.204:8001/v1',
        history=False,
        history_max_turns=50,
        history_clear_method='pop',
        temperature=0.9,
    )
    role_prompt = '你是一位汉字专家。你需要找出user提供给你的文本中的所有错别字，并给出修改意见。'
    example_prompts = [
        '例如，user发送给你的文字中有单个字的笔误："你是我的好彭友，我们明天粗去玩吧？"，你要指出"彭"应为"朋"、"粗"应为"出"。',
        '例如，user发送给你的文字中有涉及语义的笔误："我们已对全社会效益已经财务效益进行了全面分析。"，你要指出"已经"应为"以及"。',
    ]
    # nagetive_example_prompt = '需要注意的是，一个词语不完整或者多余，并不属于错别字，例如"社会效益最大"应为"社会效益最大化"、"电影院"应为"电影"就不属于错别字，不要将这种情况误判为错别字。'
    nagetive_example_prompt = ''
    style_prompt = '你的错别字修改意见要以json格式返回，具体的json格式要求是，有错别字时像这样：{"result":"有错别字", [{"原有词语":"彭友", "修改意见":"朋友"}, {"原有词语":"粗去", "修改意见":"出去"}]}，没用错别字时像这样：{"result":"无错别字"}。'
    other_requirement = '直接返回json意见，不作任何解释。一步一步想清楚。'
    llm.set_role_prompt(role_prompt+''.join(example_prompts)+nagetive_example_prompt+style_prompt+other_requirement)

    # llm.print_history()
    # llm.ask_prepare('现在帮我在这段文字中找错别字："报告所提投资优化分析适用于新型电力系统、综合能源项目、微电网项目以及传统电力系统项目，具体支持冷热电气各类机组和设备模型，负荷类型支持城市类型、工业类型等，优化目标支持社会效益最大和财务效益最佳等。计算中已经内置了8760h负荷特性、新能源出力特性已经分时电价等信息。"').get_answer_and_sync_print()

    doc = LLM_Doc('d:/server/life-agent/tools/doc/错别字案例.docx')
    doc.win32com_init()
    res_list = doc.check_wrong_written_in_docx_file(llm)
    doc.win32_close_file()

    print("*"*30+"纠错结果汇编"+"*"*30)
    jj=0
    for item in res_list:
        jj += 1
        print(f'第{jj}段检查结果：\n {item}')

def main_toc():
    llm = LLM_Qwen(
        url='http://116.62.63.204:8001/v1',
        history=False,
        history_max_turns=50,
        history_clear_method='pop',
        temperature=0.9,
    )
    file = 'd:/server/life-agent/tools/doc/南麂岛离网型微网示范工程-总报告.docx'
    import docx
    doc = docx.Document(file)
    for para in doc.paragraphs:
        if para.style.name=="Heading 1":
            print(para.text)
        if para.style.name=="Heading 2":
            print('\t'+para.text)
        if para.style.name=="Heading 3":
            print('\t\t'+para.text)
        # if para.style.name=="Heading 4":
        #     print('\t\t\t'+para.text)
    # doc = LLM_Doc(file)
    # doc.win32com_init()
    # for para in doc.get_paras():
    #     print(para)
    # doc.win32_close_file()

def main_image():
    file = 'd:/server/life-agent/tools/doc/南麂岛离网型微网示范工程-总报告.docx'
    doc = LLM_Doc(file)
    # for para in doc.get_paras():
    #     print(para)
    # for image in doc.doc.inline_shapes:
    #     print(f'image: {image.type}')
        # print(f'image: {docx.enum.shape.WD_INLINE_SHAPE(3)}')
        # print(docx.enum.shape.WD_INLINE_SHAPE.PICTURE)

    # for image in doc.get_all_inline_images():
    #     print(', '.join([
    #         f'image name: {image["name"]}',
    #         f'image size: {len(image["data"])}',
    #         f'image width: {image["width"]}',
    #         f'image height: {image["height"]}'
    #     ]))

    doc.parse_all_docx()
    # doc.print_doc_root()
    # doc.print_doc_root('2.1.7')
    # node = doc.find_doc_root('9')
    node = doc.find_doc_root('2.1.6.3')
    # node = doc.find_doc_root('2.1.7')
    doc.print_from_doc_node(node)

    # text_list = []
    # doc.get_text_from_doc_node(text_list, node)
    # print('\n'.join(text_list))


    # print(doc.get_toc_list_json_string(in_max_level=3))
    # print(doc.get_toc_json_string(in_max_level=3))

def ask_docx(in_filename='d:/server/life-agent/tools/doc/南麂岛离网型微网示范工程-总报告.docx'):
    llm = LLM_Qwen(
        history=False,
        # history_max_turns=50,
        # history_clear_method='pop',
        temperature=0.7,
        url='http://127.0.0.1:8001/v1',
        need_print=False,
    )

    file = in_filename
    doc = LLM_Doc(file, llm)
    doc.parse_all_docx()
    toc = doc.get_toc_list_json_string(in_max_level=3)



    while True:
        query = input('User: ')

        json_example = '{"head":"1.1.3"}'
        prompt = '''
        这是文档的目录结构"{toc}",
        请问这个问题"{query}"涉及的内容应该在具体的哪个章节中，不解释，请直接以"章节编号"形式返回。
        '''

        prompt = prompt.format(toc=toc, query=query, json_example=json_example)
        # print(f'--------发给LLM的prompt----------')
        # print(prompt)
        # print(f'--------发给LLM的prompt----------')
        res = llm.ask_prepare(prompt).get_answer_and_sync_print()
        print(f'Bot: {res}')
        re_result = re.search(r"\d+(.\d+)*",res).group(0)
        print(f'RE: {re_result}')

        node = doc.find_doc_root(re_result)
        # text_got = node.node_data.text
        inout_text = []
        doc.get_text_from_doc_node(inout_text, node)
        text_got = '\n'.join(inout_text)
        print(f'text_got: {text_got}')


        prompt2 = '''
        请根据材料"{text_got}"中的内容, 回答问题"{query}"。
        '''

        prompt2 = prompt2.format(text_got=text_got, query=query)
        llm.need_print = True
        res = llm.ask_prepare(prompt2).get_answer_and_sync_print()

# Color枚举类
class Color(Enum):
    red=auto()
    green=auto()
    blue=auto()

def main_table():
    # doc1 = Document('d:/server/life-agent/tools/doc/南麂岛离网型微网示范工程-总报告.docx')
    doc = LLM_Doc('d:/server/life-agent/tools/doc/南麂岛离网型微网示范工程-总报告.docx')

    # for table in doc1.tables:
    #     print('-------------------输出table-------------------')
    #     for row in table.rows:
    #         for cell in row.cells:
    #             print(cell.text, end='\t', flush=True)
    #         print()

    for block in doc.iter_block_items(doc.doc):
        # block.style.name可以直接返回：heading 1、normal、normal table
        if block.style.name == 'Normal Table':
            table = block
            print('-------------------输出table-------------------')
            for row in table.rows:
                for cell in row.cells:
                    print(cell.text, end='\t', flush=True)
                print()

    # doc.parse_all_docx()
    # # doc.print_doc_root()
    # # doc.print_doc_root('2.1.7')
    # node = doc.find_doc_root('2.1.3.2')
    # doc.print_from_doc_node(node)


def main():


    doc = LLM_Doc(in_file_name='d:/server/life-agent/tools/doc/南麂岛离网型微网示范工程-总报告.docx')
    while True:
        query = input("User: ")
        print('Assistant: ')
        gen = doc.ask_docx(query)
        for chunk in gen:
            print(chunk, end='', flush=True)
        print()

if __name__ == "__main__":
    # main_table()
    # (? <= \s)\d + (?=\s)
    main_image()
    # match_unit = re.search(r'\s*(?<=注[:：]).*', ' 注：由于各配电变压器均只有一条631或632线路支线作为进线，各支线的潮流与流入配电变压器的功率相等，简化起见，在表中有功、无功和视在功率统一表示。')
    # match_unit = re.search(r'^(\s(表|附表))\d+(.\d+)*', '\n附表16                                 南麂岛智能用电系统安装工程费用表                                 单位：元')
    # match_unit = re.search(r'(?<=单位[：|:])\s*[\u4e00-\u9fa5]+', '附表21   南麂岛10kV线路总概算表   单位： 万元')
    # match_unit = match_unit.group(0) if match_unit else ''
    # print(match_unit)

    # doc = fitz.open("D:/server/life-agent/WorldEnergyOutlook2023.pdf")
    # # 获取Document 文档对象的属性和方法
    # # 1、获取pdf 页数
    # pageCount = doc.page_count
    # print("pdf 页数", pageCount)
    #
    # # 2、获取pdf 元数据
    # metaData = doc.metadata
    # print("pdf 元数据:", metaData)

    # 3、获取pdf 目录信息
    # llm = LLM_Qwen(
    #     history=False,
    #     # history_max_turns=50,
    #     # history_clear_method='pop',
    #     temperature=0.7,
    #     url='http://127.0.0.1:8001/v1',
    #     need_print=False,
    # )

    # toc = doc.get_toc()
    # print("pdf 目录：")

    # 4、遍历para
    # for page in doc.pages(100, 102):
    #     print(f'--------{page}---------')
    #     print(page.get_text('text'))
    #     # print(f'--------{page}---------')
    #     # print(page.get_text('blocks'))
    #     print(f'--------{page}---------')
    #     # print(json.dumps(page.get_text('dict'), indent=2))
    #     text_dict = page.get_text('dict')
    #     blocks = text_dict['blocks']
    #     for block in blocks:
    #         lines = block['lines']
    #         for line in lines:
    #             spans = line['spans']
    #             for span in spans:
    #                 text = span['text']
    #                 print(f'【line】: {text}')

        # print(f'--------{page}---------')
        # print(page.get_text().encode('utf8'))
        # print(f'--------{page}---------')
        # print(page.get_text())

    # for item in toc:
    #     level = item[0]
    #     head = item[1]
    #     print(f'{"-"*level}{head}')
    #
    # # prompt = f'"{toc_json}"为一本书的目录结构列表，注意列表中每一个元素的数据结构为[level, toc_head, page]，请只把level为1和2的目录标题翻译为中文后返回给我'
    #
    #     prompt = f'把{head}翻译为中文'
    #     res = llm.ask_prepare(prompt, in_max_tokens=4096).get_answer_and_sync_print()
    #     print(res)




    # print(f'color: {Color(1)}')
    # print(f'color: {Color.blue}')


# 若win32com打开word文件报错：AttributeError: module 'win32com.gen_py.00020905-0000-0000-C000-000000000046x0x8x7' has no attribute 'CLSIDToClassMap'
# 则删除目录C:\Users\tutu\AppData\Local\Temp\gen_py\3.10中的对应缓存文件夹00020905-0000-0000-C000-000000000046x0x8x7即可

# doc.win32_close_file()若报错：pywintypes.com_error: (-2147352567, '发生意外。', (0, 'Microsoft Word', '类型不匹配', 'wdmain11.chm', 36986, -2146824070), None)
# 很可能是和wps有关，据说卸载word，win32.gencache.EnsureDispatch('Word.Application')会成功调用wps