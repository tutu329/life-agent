#!/usr/bin/env python3
# -*- coding: utf-8 -*-

import re
import pathlib
import shutil
import subprocess
from typing import List, Dict, Tuple, Optional, Union
from docx import Document
from docx.shared import Inches
import os
import argparse

class DocxOutlineExtractor:
    """
    DOCX/DOC文档大纲提取器
    支持 .doc 自动转换为 .docx（依赖 LibreOffice CLI）
    通过正则表达式系统性分析文档结构，提取章节大纲
    """

    def __init__(self):
        # 中文数字映射
        self.chinese_numbers = {
            '一': 1, '二': 2, '三': 3, '四': 4, '五': 5, '六': 6, '七': 7, '八': 8, '九': 9, '十': 10,
            '十一': 11, '十二': 12, '十三': 13, '十四': 14, '十五': 15, '十六': 16, '十七': 17, '十八': 18, '十九': 19,
            '二十': 20
        }

        # 各种章节标记的正则表达式
        self.patterns = {
            'numeric': r'^(\d+(?:\.\d+)*)\s*(.+)$',  # 1, 1.1, 1.1.1
            'chapter': r'^(?:第([一二三四五六七八九十]+)章|第(\d+)章)\s*(.+)$',  # 第一章, 第1章
            'chinese_item': r'^([一二三四五六七八九十]+)、\s*(.+)$',  # 一、二、
            'numbered_item': r'^(\d+)[.、]\s*(.+)$',  # 1. 或 1、
            'parenthesis': r'^(\d+)[)）]\s*(.+)$',  # 1) 或 1）
            'bracket': r'^[(\（](\d+)[)\）]\s*(.+)$',  # (1) 或 （1）
            'letter': r'^([a-z])[)\)]\s*(.+)$',  # a) 或 b)
            'bracket_letter': r'^[(\（]([a-z])[)\）]\s*(.+)$',  # (a) 或 (b)
            'date': r'^\d{4}\s*年\s*\d{1,2}\s*月\s*\d{1,2}\s*日',  # 过滤日期
            'hash_date': r'^#\s*\d{4}\s*年\s*\d{1,2}\s*月\s*\d{1,2}\s*日',  # 过滤# 2024年12月16日
        }

    def _convert_with_libreoffice(self, src: pathlib.Path) -> pathlib.Path:
        """
        使用 LibreOffice CLI 将 .doc 转为 .docx。

        LibreOffice 会把输出文件放到工作目录（缺省是 cwd）。
        若目标存在则覆盖。
        """
        soffice = shutil.which("soffice") or shutil.which("libreoffice")
        if not soffice:
            raise RuntimeError("未找到 LibreOffice，可执行文件 soffice / libreoffice 不在 PATH。\n"
                               "请安装 LibreOffice: sudo apt install libreoffice")

        dst = src.with_suffix(".docx")
        # --headless: 无界面；--convert-to: 格式；--outdir: 指定输出目录
        cmd = [
            soffice,
            "--headless",
            "--convert-to",
            "docx",
            "--outdir",
            str(src.parent),
            str(src),
        ]
        proc = subprocess.run(cmd, stdout=subprocess.PIPE, stderr=subprocess.PIPE)
        if proc.returncode:
            raise RuntimeError(
                f"LibreOffice 转换失败 (exit {proc.returncode}).\nSTDERR:\n{proc.stderr.decode(errors='ignore')}"
            )
        if not dst.exists():
            raise RuntimeError("LibreOffice 返回 0，但未找到生成的 .docx 文件。")
        return dst

    def _ensure_docx(self, path: Union[str, pathlib.Path]) -> str:
        """若输入是 .doc，调用 LibreOffice 转换；否则直接返回 .docx 路径。"""
        p = pathlib.Path(path)
        if p.suffix.lower() == ".docx":
            return str(p)
        if p.suffix.lower() == ".doc":
            print(f"检测到 .doc 文件，正在转换为 .docx: {p}")
            converted_path = self._convert_with_libreoffice(p)
            print(f"转换完成: {converted_path}")
            return str(converted_path)
        raise ValueError("仅支持 .doc / .docx 文件")

    def extract_outline(self, file_path: str, max_depth: int = 6) -> List[Dict]:
        """
        提取文档大纲

        Args:
            file_path: doc/docx文件路径
            max_depth: 最大级别深度

        Returns:
            章节列表，每个章节包含：title, level, number
        """
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"文件不存在: {file_path}")

        # 确保是 .docx 格式（如果是 .doc 则自动转换）
        docx_path = self._ensure_docx(file_path)

        # 读取文档
        doc = Document(docx_path)
        paragraphs = []

        # 提取所有段落
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                paragraphs.append({
                    'text': text,
                    'style': para.style.name if para.style else None,
                    'level': self._get_outline_level(para)
                })

        # 分析章节体系
        system_type = self._analyze_numbering_system(paragraphs)

        # 提取章节
        chapters = self._extract_chapters(paragraphs, system_type, max_depth)

        return chapters

    def _get_outline_level(self, paragraph) -> Optional[int]:
        """获取段落的大纲级别"""
        try:
            if paragraph.style and hasattr(paragraph.style, 'base_style'):
                style_name = paragraph.style.tool_name
                if 'Heading' in style_name:
                    # 提取标题级别
                    level_match = re.search(r'(\d+)', style_name)
                    if level_match:
                        return int(level_match.group(1))
            return None
        except:
            return None

    def _analyze_numbering_system(self, paragraphs: List[Dict]) -> str:
        """
        分析文档使用的编号系统

        Returns:
            'numeric': 数字体系 (1, 1.1, 1.1.1)
            'mixed': 混合体系 (第一章, 一、, 1、, 1), (1), a))
        """
        numeric_pattern = r'^\d+(?:\.\d+)+\s'  # 检查多级数字编号
        single_numeric_pattern = r'^\d+\s+[^\d、)）]'  # 检查单级数字编号，如 "1 总论"
        numeric_count = 0

        chapter_pattern = r'^(?:第[一二三四五六七八九十\d]+章)'
        chinese_item_pattern = r'^[一二三四五六七八九十]+、'
        mixed_indicators = 0

        for para in paragraphs[:50]:  # 只检查前50个段落
            text = para['text']

            # 检查多级数字编号
            if re.match(numeric_pattern, text):
                numeric_count += 1

            # 检查单级数字编号
            if re.match(single_numeric_pattern, text):
                numeric_count += 1

            # 检查混合体系指示器
            if re.match(chapter_pattern, text) or re.match(chinese_item_pattern, text):
                mixed_indicators += 1

        # 如果有较多的数字编号，认为是数字体系
        if numeric_count >= 3:
            return 'numeric'
        else:
            return 'mixed'

    def _extract_chapters(self, paragraphs: List[Dict], system_type: str, max_depth: int) -> List[Dict]:
        """提取章节"""
        chapters = []

        for para in paragraphs:
            text = para['text']

            # 过滤日期和非章节内容
            if self._is_filtered_content(text):
                continue

            # 尝试正则匹配
            chapter_info = self._match_chapter_pattern(text, system_type)

            # 如果正则匹配失败，尝试使用大纲级别
            if not chapter_info and para['level']:
                chapter_info = self._extract_by_outline_level(text, para['level'])

            if chapter_info and chapter_info['level'] <= max_depth:
                chapters.append(chapter_info)

        return self._filter_toc_content(chapters)

    def _is_filtered_content(self, text: str) -> bool:
        """过滤不需要的内容"""
        # 过滤日期
        if re.match(self.patterns['date'], text) or re.match(self.patterns['hash_date'], text):
            return True

        # 过滤目录标识
        if re.match(r'^目\s*录', text) or re.match(r'^contents?$', text.lower()):
            return True

        # 过滤页码行
        if re.match(r'^\d+\s*$', text):
            return True

        # 过滤制表符行（目录中的页码对齐）
        if '\t' in text and re.match(r'^.+\t+\d+\s*$', text):
            return True

        # 过滤目录内容：章节标题后跟多个空格和页码
        if re.match(r'^.+\s{3,}\d+\s*$', text):
            return True

        # 过滤包含详细描述的工具列表项
        if re.match(r'^\d+\s*、.*[:：]', text):
            return True

        # 过滤表格标题行
        if re.match(r'.*表\s*$', text) and len(text) < 20:
            return True

        return False

    def _match_chapter_pattern(self, text: str, system_type: str) -> Optional[Dict]:
        """匹配章节模式"""

        if system_type == 'numeric':
            # 数字体系：优先匹配多级数字编号
            match = re.match(self.patterns['numeric'], text)
            if match:
                number = match.group(1)
                title = match.group(2).strip()
                level = len(number.split('.'))
                return {'title': title, 'level': level, 'number': number}

            # 检查单级数字编号，如 "1 总论"
            single_match = re.match(r'^(\d+)\s+([^\d、)）].*)$', text)
            if single_match:
                number = single_match.group(1)
                title = single_match.group(2).strip()
                # 过滤掉不是真正章节标题的内容
                if self._is_valid_chapter_title(title):
                    return {'title': title, 'level': 1, 'number': number}

        else:  # mixed system
            # 混合体系：按级别顺序匹配

            # 第一级：第一章, 第1章
            match = re.match(self.patterns['chapter'], text)
            if match:
                chinese_num = match.group(1)
                arabic_num = match.group(2)
                title = match.group(3).strip()
                if chinese_num:
                    number = chinese_num
                else:
                    number = arabic_num
                return {'title': title, 'level': 1, 'number': f"第{number}章"}

            # 第二级：一、二、
            match = re.match(self.patterns['chinese_item'], text)
            if match:
                chinese_num = match.group(1)
                title = match.group(2).strip()
                return {'title': title, 'level': 2, 'number': chinese_num}

            # 第三级：1、1.
            match = re.match(self.patterns['numbered_item'], text)
            if match:
                number = match.group(1)
                title = match.group(2).strip()
                # 在混合体系中，检查是否为最高级别的数字
                if self._is_top_level_number(text, number):
                    return {'title': title, 'level': 1, 'number': number}
                else:
                    return {'title': title, 'level': 3, 'number': number}

            # 第四级：1) 1）
            match = re.match(self.patterns['parenthesis'], text)
            if match:
                number = match.group(1)
                title = match.group(2).strip()
                return {'title': title, 'level': 4, 'number': f"{number})"}

            # 第五级：(1) （1）
            match = re.match(self.patterns['bracket'], text)
            if match:
                number = match.group(1)
                title = match.group(2).strip()
                return {'title': title, 'level': 5, 'number': f"({number})"}

            # 第六级：a) (a)
            match = re.match(self.patterns['letter'], text)
            if match:
                letter = match.group(1)
                title = match.group(2).strip()
                return {'title': title, 'level': 6, 'number': f"{letter})"}

            match = re.match(self.patterns['bracket_letter'], text)
            if match:
                letter = match.group(1)
                title = match.group(2).strip()
                return {'title': title, 'level': 6, 'number': f"({letter})"}

        return None

    def _is_top_level_number(self, text: str, number: str) -> bool:
        """
        在混合体系中判断数字是否为最高级别
        如果文档中没有"第一章"、"一、"等更高级别的标记，则"1、"为最高级别
        """
        # 这里需要根据实际文档分析，简化处理
        # 实际应用中可能需要更复杂的逻辑
        return False

    def _is_valid_chapter_title(self, title: str) -> bool:
        """
        判断是否是有效的章节标题
        过滤掉工具列表、详细描述等非章节标题内容
        """
        # 过滤掉包含冒号的详细描述
        if '：' in title or ':' in title:
            return False

        # 过滤掉以"、"开头的内容（通常是列表项）
        if title.startswith('、'):
            return False

        # 过滤掉过长的标题（可能是详细描述）
        if len(title) > 50:
            return False

        # 过滤掉包含特定关键词的内容
        filter_keywords = ['工具', '实现', '提供', '支持', '功能', '能力', '处理', '自动', '批量']
        for keyword in filter_keywords:
            if keyword in title and len(title) > 20:
                return False

        return True

    def _extract_by_outline_level(self, text: str, level: int) -> Optional[Dict]:
        """根据大纲级别提取章节"""
        # 处理无编号的章节标题，如"附录"、"附表"等
        if level <= 6:
            return {'title': text, 'level': level, 'number': ''}
        return None

    def _filter_toc_content(self, chapters: List[Dict]) -> List[Dict]:
        """过滤目录内容"""
        filtered_chapters = []

        for i, chapter in enumerate(chapters):
            # 如果连续出现多个相同级别的章节，且标题很短，可能是目录
            if self._is_likely_toc_entry(chapter, chapters, i):
                continue

            filtered_chapters.append(chapter)

        return filtered_chapters

    def _is_likely_toc_entry(self, chapter: Dict, all_chapters: List[Dict], index: int) -> bool:
        """判断是否可能是目录条目"""
        # 简单的启发式规则：
        # 1. 标题很短（少于2个字符）且在文档前部出现
        # 2. 后面紧跟着数字（可能是页码）

        title = chapter['title']

        # 标题太短可能是目录（但"总论"、"概述"等2字标题是有效的）
        if len(title) < 2 and index < len(all_chapters) * 0.1:
            return True

        # 如果标题是单个字符且在前10%的位置，可能是目录
        if len(title) == 1 and index < len(all_chapters) * 0.1:
            return True

        return False

    def format_outline(self, chapters: List[Dict]) -> str:
        """格式化大纲输出"""
        result = []

        for chapter in chapters:
            level = chapter['level']
            title = chapter['title']
            number = chapter['number']

            # 生成缩进
            indent = '#' * level

            if number:
                result.append(f"{indent} {number} {title}")
            else:
                result.append(f"{indent} {title}")

        return '\n'.join(result)


def extract_file_outline(file_path: str, max_depth: int = 6, show_details: bool = False):
    """提取文件大纲"""
    extractor = DocxOutlineExtractor()

    if not os.path.exists(file_path):
        print(f"错误: 文件不存在: {file_path}")
        return

    # 检查文件类型
    file_ext = pathlib.Path(file_path).suffix.lower()
    if file_ext not in ['.doc', '.docx']:
        print(f"错误: 不支持的文件格式 {file_ext}，仅支持 .doc 和 .docx 文件")
        return

    print(f"正在分析文件: {file_path}")

    try:
        chapters = extractor.extract_outline(file_path, max_depth=max_depth)

        print(f"\n提取到 {len(chapters)} 个章节:")
        print("=" * 50)
        print(extractor.format_outline(chapters))

        if show_details:
            # 显示详细信息
            print("\n详细信息:")
            print("=" * 50)
            for i, chapter in enumerate(chapters):
                print(f"{i + 1}. 级别{chapter['level']} - 编号: '{chapter['number']}' - 标题: '{chapter['title']}'")

    except Exception as e:
        print(f"分析失败: {e}")
        import traceback
        traceback.print_exc()


def main():
    """主函数，处理命令行参数"""
    parser = argparse.ArgumentParser(
        description="从 DOC/DOCX 文档中提取大纲结构",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
使用示例:
  python from_server_docx_outline.py 我的模板.docx
  python from_server_docx_outline.py 报告.doc --max-depth 4
  python from_server_docx_outline.py 文档.docx --details

依赖要求:
  - 处理 .doc 文件需要安装 LibreOffice:
    sudo apt install libreoffice
  - 处理 .docx 文件仅需要 python-docx 库
        """
    )

    parser.add_argument(
        'file_path',
        help='要分析的 DOC/DOCX 文件路径'
    )

    parser.add_argument(
        '--max-depth',
        type=int,
        default=6,
        help='最大级别深度 (默认: 6)'
    )

    parser.add_argument(
        '--details',
        action='store_true',
        help='显示详细信息'
    )

    args = parser.parse_args()

    # 执行大纲提取
    extract_file_outline(args.file_path, args.max_depth, args.details)


if __name__ == "__main__":
    main()

# #!/usr/bin/env python3
# # -*- coding: utf-8 -*-
#
# # sudo apt update
# # sudo apt install libreoffice       # LibreOffice CLI，~300 MB
# # grep -qxF 'export LD_LIBRARY_PATH=/usr/lib/libreoffice/program:$LD_LIBRARY_PATH' ~/.bashrc || echo 'export LD_LIBRARY_PATH=/usr/lib/libreoffice/program:$LD_LIBRARY_PATH' >> ~/.bashrc && source ~/.bashrc
#
# """
# Doc / Docx outline extractor —— **LibreOffice 方案**
# ----------------------------------------------------
# *   支持 .doc 自动转换为 .docx，依赖 **LibreOffice CLI (soffice)**。
# *   直接处理 .docx 文件无需转换。
# *   可按层级输出 Markdown 或 JSON 目录树。
#
# """
# from __future__ import annotations
#
# import argparse
# import json
# import pathlib
# import re
# import shutil
# import subprocess
# import sys
# from dataclasses import dataclass, field
# from typing import Any, List, Optional, Union
#
# from docx import Document
# from docx.text.paragraph import Paragraph
#
# # ---------------------------------------------------------------------------
# # .doc → .docx 自动转换（LibreOffice）
# # ---------------------------------------------------------------------------
#
#
# def _convert_with_libreoffice(src: pathlib.Path) -> pathlib.Path:
#     """
#     使用 LibreOffice CLI 将 .doc 转为 .docx。
#
#     *   LibreOffice 会把输出文件放到工作目录（缺省是 cwd）。
#     *   若目标存在则覆盖。
#     """
#     soffice = shutil.which("soffice") or shutil.which("libreoffice")
#     if not soffice:
#         raise RuntimeError("未找到 LibreOffice，可执行文件 soffice / libreoffice 不在 PATH。")
#
#     dst = src.with_suffix(".docx")
#     # --headless: 无界面；--convert-to: 格式；--outdir: 指定输出目录
#     cmd = [
#         soffice,
#         "--headless",
#         "--convert-to",
#         "docx",
#         "--outdir",
#         str(src.parent),
#         str(src),
#     ]
#     proc = subprocess.run(cmd, stdout=subprocess.PIPE, stderr=subprocess.PIPE)
#     if proc.returncode:
#         raise RuntimeError(
#             f"LibreOffice 转换失败 (exit {proc.returncode}).\nSTDERR:\n{proc.stderr.decode(errors='ignore')}"
#         )
#     if not dst.exists():
#         raise RuntimeError("LibreOffice 返回 0，但未找到生成的 .docx 文件。")
#     return dst
#
#
# def ensure_docx(path: Union[str, pathlib.Path]) -> str:
#     """若输入是 .doc，调用 LibreOffice 转换；否则直接返回 .docx 路径。"""
#     p = pathlib.Path(path)
#     if p.suffix.lower() == ".docx":
#         return str(p)
#     if p.suffix.lower() != ".doc":
#         raise ValueError("仅支持 .doc / .docx 文件")
#     return str(_convert_with_libreoffice(p))
#
#
# # ---------------------------------------------------------------------------
# # 数据结构
# # ---------------------------------------------------------------------------
#
#
# @dataclass
# class OutlineNode:
#     level: int
#     title: str
#     paragraph_index: int
#     children: List["OutlineNode"] = field(default_factory=list)
#
#     def add_child(self, node: "OutlineNode") -> None:
#         self.children.append(node)
#
#     def to_dict(self) -> dict:  # 用于 JSON 序列化
#         return {
#             "level": self.level,
#             "title": self.title,
#             "paragraph_index": self.paragraph_index,
#             "children": [c.to_dict() for c in self.children],
#         }
#
#     # -------- Markdown helpers --------
#     def _collect_md(self, out: List[str], max_hash: int) -> None:
#         if self.level > 0:
#             out.append(f"{'#' * min(self.level, max_hash)} {self.title}")
#         for c in self.children:
#             c._collect_md(out, max_hash)
#
#     def to_markdown(self, max_hash: int = 6) -> str:
#         out: List[str] = []
#         self._collect_md(out, max_hash)
#         return "\n".join(out)
#
#
# @dataclass
# class OutlineTree:
#     roots: List[OutlineNode]
#
#     def to_markdown(self) -> str:
#         return "\n".join(r.to_markdown() for r in self.roots)
#
#     def to_json(self, **json_kwargs: Any) -> str:
#         return json.dumps([r.to_dict() for r in self.roots], **json_kwargs)
#
#
# # ---------------------------------------------------------------------------
# # 核心提纲提取器（与你之前版本相同）
# # ---------------------------------------------------------------------------
#
#
# class DocxOutlineExtractor:
#     """支持中英混合编号的提纲提取器，可自动忽略目录页。"""
#
#     _DELIMS = " 、.．)）]>»›››·‧•\t --–—\u3000"
#
#     _TOKEN_RE = re.compile(
#         r"""^\s*                 # lead space
#         [\(（\[]?                 # optional opening bracket
#         (                         # token
#             (?:[0-9]+(?:\.[0-9]+)*) |   # 1 / 1.1 / 1.1.1 …
#             (?:[IVXLCDM]+)        |       # Roman upper
#             (?:[ivxlcdm]+)        |       # Roman lower
#             (?:[一二三四五六七八九十百千]+) | # Chinese numeral
#             (?:[a-z])             |       # latin lower
#             (?:[A-Z])                     # latin upper
#         )
#         [\)）\]]?                # optional closing bracket INSIDE token
#         """,
#         re.VERBOSE,
#     )
#
#     _TOC_HEURISTIC_RE = re.compile(r"(\s{2,}|\.{2,})\s*\d+\s*$")
#     _CHINESE_RE = re.compile(r"^[一二三四五六七八九十百千]+$")
#
#     def __init__(self, file_path: str, max_level: Optional[int] = 4):
#         real_path = ensure_docx(file_path)
#         self.doc = Document(real_path)
#         self.paragraphs = self.doc.paragraphs
#         self.max_level = max_level
#         self.has_dot_chain = self._detect_dot_depth() > 0
#         self.max_dot_chain = self._detect_dot_depth()
#
#     # --------------- Public API ---------------
#     def outline(self) -> OutlineTree:
#         if self._has_heading():
#             return self._outline_from_heading()
#         if self._has_numbering():
#             return self._outline_from_numbering()
#         return self._outline_from_plain()
#
#     # --------------- Internal helpers ---------------
#     def _is_toc_entry(self, p: Paragraph) -> bool:
#         style_name = getattr(p.style, "name", "").lower()
#         if "toc" in style_name or "目录" in style_name:
#             return True
#         text = p.text.strip()
#         if not text:
#             return False
#         if "\t" in text and text.split("\t")[-1].strip().isdigit():
#             return True
#         if self._TOC_HEURISTIC_RE.search(text):
#             return True
#         return False
#
#     def _outline_from_heading(self) -> OutlineTree:
#         roots, stack = [], []
#         for i, p in enumerate(self.paragraphs):
#             if self._is_toc_entry(p):
#                 continue
#             style = getattr(p.style, "name", "")
#             if not style.startswith("Heading"):
#                 continue
#             try:
#                 lvl = int(style.split()[1])
#             except Exception:
#                 lvl = 1
#             if self._skip(lvl):
#                 continue
#             self._insert(OutlineNode(lvl, p.text.strip(), i), stack, roots)
#         return OutlineTree(roots)
#
#     def _outline_from_numbering(self) -> OutlineTree:
#         roots, stack = [], []
#         for i, p in enumerate(self.paragraphs):
#             if self._is_toc_entry(p):
#                 continue
#             numPr = p._p.pPr.numPr if p._p.pPr is not None else None
#             if numPr is None:
#                 continue
#             try:
#                 lvl = int(numPr.ilvl.val) + 1
#             except Exception:
#                 lvl = 1
#             if self._skip(lvl):
#                 continue
#             self._insert(OutlineNode(lvl, p.text.strip(), i), stack, roots)
#         return OutlineTree(roots)
#
#     def _outline_from_plain(self) -> OutlineTree:
#         roots, stack = [], []
#         for i, p in enumerate(self.paragraphs):
#             if self._is_toc_entry(p):
#                 continue
#             txt = p.text.rstrip()
#             if not txt:
#                 continue
#             m = self._TOKEN_RE.match(txt)
#             if not m:
#                 continue
#             end = m.end()
#             delim = txt[end] if end < len(txt) else ""
#             token_includes_bracket = m.group(0).strip().endswith((")", "）", "]"))
#             if end < len(txt) and not token_includes_bracket and delim not in self._DELIMS:
#                 continue
#             raw = m.group(0)
#             lvl = self._infer_level(raw)
#             if self._skip(lvl):
#                 continue
#             self._insert(OutlineNode(lvl, txt.strip(), i), stack, roots)
#         return OutlineTree(roots)
#
#     # ---- level inference & util ----
#     def _detect_dot_depth(self) -> int:
#         max_depth = 0
#         for p in self.paragraphs:
#             if self._is_toc_entry(p):
#                 continue
#             m = self._TOKEN_RE.match(p.text.lstrip())
#             if m and "." in m.group(0):
#                 token = (
#                     m.group(0)
#                     .lstrip(" (（[")
#                     .rstrip(self._DELIMS)
#                     .rstrip(" )）]")
#                 )
#                 if token.replace(".", "").isdigit():
#                     max_depth = max(max_depth, token.count(".") + 1)
#         return max_depth
#
#     def _infer_level(self, raw: str) -> int:
#         token = raw.strip().lstrip(" (（[").rstrip(self._DELIMS).rstrip(" )）]")
#         if "." in token and token.replace(".", "").isdigit():
#             return token.count(".") + 1
#         is_chinese = bool(self._CHINESE_RE.fullmatch(token))
#         is_digit = token.isdigit()
#         is_letter = len(token) == 1 and token.isalpha()
#         has_paren = any(c in raw for c in ")）")
#         has_comma = any(c in raw for c in "、.．")
#
#         if self.has_dot_chain:
#             if is_digit and not has_comma and not has_paren:
#                 return 1
#             base = self.max_dot_chain
#             if is_chinese:
#                 return base + 1
#             if is_digit and has_comma:
#                 return base + 2
#             if is_digit and has_paren:
#                 return base + 3
#             if is_letter and has_paren:
#                 return base + 4
#             return 99
#         else:
#             if is_chinese:
#                 return 1
#             if is_digit and has_comma:
#                 return 2
#             if is_digit and has_paren:
#                 return 3
#             if is_letter and has_paren:
#                 return 4
#             if is_digit:
#                 return 2
#             return 99
#
#     def _insert(self, node: OutlineNode, stack: List[OutlineNode], roots: List[OutlineNode]):
#         while len(stack) >= node.level:
#             stack.pop()
#         (stack[-1].add_child(node) if stack else roots.append(node))
#         stack.append(node)
#
#     def _skip(self, lvl: int) -> bool:
#         return self.max_level is not None and lvl > self.max_level
#
#     def _has_heading(self) -> bool:
#         for p in self.paragraphs:
#             style = getattr(p.style, "name", "")
#             if "toc" in style.lower() or "目录" in style:
#                 continue
#             if style.startswith("Heading"):
#                 return True
#         return False
#
#     def _has_numbering(self) -> bool:
#         return any(
#             p._p.pPr is not None and p._p.pPr.numPr is not None and not self._is_toc_entry(p)
#             for p in self.paragraphs
#         )
#
#
# # ---------------------------------------------------------------------------
# # CLI
# # ---------------------------------------------------------------------------
#
#
# def _main() -> None:
#     parser = argparse.ArgumentParser(
#         description="Extract outline from DOC / DOCX (LibreOffice 方案)",
#     )
#     parser.add_argument("file", help="Path to .doc 或 .docx 文件")
#     parser.add_argument("-f", "--format", choices=["md", "json"], default="md", help="输出格式 (md|json)")
#     parser.add_argument("-l", "--max-level", type=int, default=4, help="最大层级 (默认 4)")
#     args = parser.parse_args()
#
#     if not pathlib.Path(args.file).exists():
#         sys.exit(f"文件不存在: {args.file}")
#
#     tree = DocxOutlineExtractor(args.file, args.max_level).outline()
#     if args.format == "md":
#         print(tree.to_markdown())
#     else:
#         print(tree.to_json(ensure_ascii=False, indent=2))
#
#
# if __name__ == "__main__":
#     _main()
#
