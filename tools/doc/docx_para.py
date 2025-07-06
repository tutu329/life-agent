#!/usr/bin/env python3
# -*- coding: utf-8 -*-

import re
import pathlib
import os
from typing import List, Dict, Optional, Tuple, Any
from docx import Document
from docx.table import Table
from docx.text.paragraph import Paragraph
import argparse

from tools.doc.docx_outline import DocxOutlineExtractor

class DocxParser:
    """
    DOCX文档章节内容解析器
    基于DocxOutlineExtractor的章节结构分析，提取指定章节的详细内容
    """

    def __init__(self, file_path: str):
        """
        初始化DocxParser

        Args:
            file_path: doc/docx文件路径
        """
        self.file_path = file_path
        self.outline_extractor = DocxOutlineExtractor()
        self.docx_path = None
        self.doc = None
        self.chapters = []
        self.paragraphs = []
        self._load_document()

    def _load_document(self):
        """加载文档并提取章节结构"""
        if not os.path.exists(self.file_path):
            raise FileNotFoundError(f"文件不存在: {self.file_path}")

        # 确保是 .docx 格式（如果是 .doc 则自动转换）
        self.docx_path = self.outline_extractor._ensure_docx(self.file_path)

        # 加载文档
        self.doc = Document(self.docx_path)

        # 获取章节结构 - 直接使用已转换的docx路径，避免重复转换
        self.chapters = self.outline_extractor.extract_outline(self.docx_path, max_depth=8)

        # 获取所有段落信息
        self._extract_paragraphs()

    def _extract_paragraphs(self):
        """提取所有段落信息，包括文本和表格"""
        self.paragraphs = []

        # 使用更简单的方法：直接遍历段落和表格
        # 为了保持顺序，我们需要通过document的body元素来获取
        for element in self.doc.element.body:
            if element.tag.endswith('p'):  # 段落
                para = None
                for p in self.doc.paragraphs:
                    if p._element == element:
                        para = p
                        break

                if para:
                    self.paragraphs.append({
                        'type': 'paragraph',
                        'element': para,
                        'text': para.text.strip(),
                        'style': para.style.name if para.style else None
                    })

            elif element.tag.endswith('tbl'):  # 表格
                table = None
                for t in self.doc.tables:
                    if t._element == element:
                        table = t
                        break

                if table:
                    self.paragraphs.append({
                        'type': 'table',
                        'element': table,
                        'text': '',
                        'style': None
                    })

    def get_chapter(self, chapter_number: str) -> str:
        """
        获取指定章节的内容

        Args:
            chapter_number: 章节编号，如 "3.2"

        Returns:
            章节内容字符串
        """
        if not self.chapters:
            return "未找到任何章节结构"

        # 查找目标章节及其子章节
        target_chapters = self._find_target_chapters(chapter_number)

        if not target_chapters:
            return f"未找到章节 {chapter_number}"

        # 提取章节内容
        content = self._extract_chapter_content(target_chapters)

        return content

    def _find_target_chapters(self, chapter_number: str) -> List[Dict]:
        """
        查找目标章节及其所有子章节

        Args:
            chapter_number: 章节编号，如 "3.2"

        Returns:
            包含目标章节及其子章节的列表
        """
        target_chapters = []

        # 检查是否是精确的章节编号匹配
        exact_match_found = False

        for chapter in self.chapters:
            number = chapter.get('number', '')

            # 精确匹配目标章节
            if self._match_chapter_number(number, chapter_number):
                target_chapters.append(chapter)
                exact_match_found = True

        # 如果找到精确匹配，且是主章节（如"1"），则只返回主章节
        # 主章节的内容范围会自动包含所有子章节内容
        if exact_match_found:
            # 检查查找的是否是主章节（不包含小数点，如"1", "2", "7"）
            if '.' not in chapter_number:
                # 对于主章节，只返回主章节，不返回子章节
                # 这样可以避免重复内容
                main_chapters = [ch for ch in target_chapters if ch.get('number', '') == chapter_number]
                if main_chapters:
                    return main_chapters

        # 如果没有找到精确匹配，或者是子章节查找，则包含子章节
        if not exact_match_found:
            for chapter in self.chapters:
                number = chapter.get('number', '')

                # 匹配子章节
                if self._is_sub_chapter(number, chapter_number):
                    target_chapters.append(chapter)

        # 按章节号排序
        target_chapters.sort(key=lambda x: self._get_sort_key(x.get('number', '')))

        return target_chapters

    def _match_chapter_number(self, chapter_num: str, target_num: str) -> bool:
        """
        匹配章节编号

        Args:
            chapter_num: 章节编号
            target_num: 目标编号

        Returns:
            是否匹配
        """
        # 提取数字部分进行比较
        chapter_digits = self._extract_number_part(chapter_num)
        target_digits = self._extract_number_part(target_num)

        return chapter_digits == target_digits

    def _is_sub_chapter(self, chapter_num: str, parent_num: str) -> bool:
        """
        判断是否为子章节

        Args:
            chapter_num: 章节编号
            parent_num: 父章节编号

        Returns:
            是否为子章节
        """
        chapter_digits = self._extract_number_part(chapter_num)
        parent_digits = self._extract_number_part(parent_num)

        # 检查是否以父章节编号开头且有更多层级
        if chapter_digits.startswith(parent_digits + '.'):
            return True

        return False

    def _extract_number_part(self, chapter_num: str) -> str:
        """
        提取章节编号的数字部分

        Args:
            chapter_num: 章节编号字符串

        Returns:
            数字部分，如 "3.2"
        """
        # 匹配数字编号模式
        match = re.search(r'(\d+(?:\.\d+)*)', chapter_num)
        if match:
            return match.group(1)

        # 如果没有数字，返回原字符串
        return chapter_num

    def _get_sort_key(self, chapter_num: str) -> Tuple:
        """
        获取章节编号的排序键

        Args:
            chapter_num: 章节编号

        Returns:
            排序键元组
        """
        digits = self._extract_number_part(chapter_num)
        if digits:
            try:
                return tuple(int(x) for x in digits.split('.'))
            except:
                pass

        return (0,)

    def _extract_chapter_content(self, target_chapters: List[Dict]) -> str:
        """
        提取章节内容

        Args:
            target_chapters: 目标章节列表

        Returns:
            格式化的章节内容
        """
        if not target_chapters:
            return ""

        # 找到每个章节标题在文档中的位置
        chapter_positions = []
        for chapter in target_chapters:
            title = chapter['title']
            number = chapter.get('number', '')

            # 查找章节标题在段落中的位置
            found = False
            for i, para_info in enumerate(self.paragraphs):
                if para_info['type'] == 'paragraph':
                    para_text = para_info['text']

                    # 尝试匹配章节标题
                    if self._is_chapter_title_match(para_text, title, number):
                        chapter_positions.append({
                            'chapter': chapter,
                            'start_pos': i,
                            'title_text': para_text
                        })
                        found = True
                        break

            # 如果没有找到，可以考虑其他匹配策略
            if not found:
                pass  # 暂时跳过

        # 如果没有找到任何章节位置，返回提示
        if not chapter_positions:
            return "未找到指定章节的内容"

        # 按位置排序
        chapter_positions.sort(key=lambda x: x['start_pos'])

        # 提取内容
        content_parts = []

        for i, pos_info in enumerate(chapter_positions):
            chapter = pos_info['chapter']
            start_pos = pos_info['start_pos']

            # 确定结束位置 - 需要找到下一个非子章节的位置
            end_pos = self._find_chapter_end_position(start_pos, chapter, chapter_positions)

            # 提取章节内容
            section_content = self._extract_section_content(
                start_pos, end_pos, chapter, pos_info['title_text']
            )

            if section_content:
                content_parts.append(section_content)

        return '\n\n'.join(content_parts)

    def _find_chapter_end_position(self, start_pos: int, current_chapter: Dict, all_positions: List[Dict]) -> int:
        """
        查找章节的结束位置

        Args:
            start_pos: 当前章节开始位置
            current_chapter: 当前章节信息
            all_positions: 所有章节位置列表

        Returns:
            章节结束位置
        """
        current_number = current_chapter.get('number', '')
        current_level = len(current_number.split('.')) if current_number else 0

        # 查找所有章节中下一个同级或上级章节的位置
        for pos_info in all_positions:
            other_pos = pos_info['start_pos']
            other_chapter = pos_info['chapter']
            other_number = other_chapter.get('number', '')

            # 跳过当前章节
            if other_pos <= start_pos:
                continue

            # 检查是否是同级或上级章节
            if other_number:
                other_level = len(other_number.split('.'))

                # 如果是同级或上级章节，这里就是结束位置
                if other_level <= current_level:
                    return other_pos

        # 如果没有找到同级或上级章节，在整个文档中查找
        return self._find_next_sibling_chapter_in_doc(start_pos, current_chapter)

    def _find_next_sibling_chapter_in_doc(self, start_pos: int, current_chapter: Dict) -> int:
        """
        在整个文档中查找下一个同级章节的位置
        """
        current_number = current_chapter.get('number', '')
        current_level = len(current_number.split('.')) if current_number else 0

        # 从当前位置开始查找
        for i in range(start_pos + 1, len(self.paragraphs)):
            para_info = self.paragraphs[i]
            if para_info['type'] != 'paragraph':
                continue

            para_text = para_info['text'].strip()
            if not para_text:
                continue

            # 检查是否匹配章节模式
            for chapter in self.chapters:
                ch_number = chapter.get('number', '')
                ch_title = chapter.get('title', '')

                if self._is_chapter_title_match(para_text, ch_title, ch_number):
                    # 检查级别
                    if ch_number:
                        ch_level = len(ch_number.split('.'))
                        if ch_level <= current_level:
                            return i

        # 如果没有找到，返回文档结尾
        return len(self.paragraphs)

    def _is_chapter_title_match(self, para_text: str, title: str, number: str) -> bool:
        """
        判断段落文本是否匹配章节标题

        Args:
            para_text: 段落文本
            title: 章节标题
            number: 章节编号

        Returns:
            是否匹配
        """
        # 跳过空段落
        if not para_text.strip():
            return False

        # 排除目录内容：检查是否包含制表符+数字（页码）
        if self._is_toc_entry(para_text):
            return False

        # 排除只包含章节编号和页码的行
        if number and re.match(rf'^{re.escape(number)}\s+\d+\s*$', para_text):
            return False

        # **优先匹配完整的章节标题模式**
        if number:
            # 尝试匹配 "编号 标题" 或 "编号标题" 的格式
            full_title_patterns = [
                f"{number} {title}",  # "4.3.3 技术架构"
                f"{number}{title}",  # "4.3.3技术架构"
                f"{number}、{title}",  # "4.3.3、技术架构"
                f"{number}．{title}",  # "4.3.3．技术架构"
            ]

            for pattern in full_title_patterns:
                if pattern in para_text:
                    # 额外检查：确保不是目录条目
                    if not self._looks_like_toc_pattern(para_text, pattern):
                        return True

        # **严格匹配：只匹配以章节编号开头的段落**
        if number and para_text.startswith(number):
            # 检查编号后是否紧跟标题或分隔符
            if (para_text.startswith(f"{number} {title}") or
                    para_text.startswith(f"{number}{title}") or
                    para_text.startswith(f"{number}、{title}") or
                    para_text.startswith(f"{number}．{title}")):
                # 确保不是目录条目
                if not self._looks_like_toc_pattern(para_text, number):
                    return True

        # **避免误匹配：不允许仅通过标题关键词匹配**
        # 这里移除了原来的 title in para_text 的宽松匹配
        # 因为这容易导致误匹配，如"技术架构"出现在参考文献中

        return False

    def _is_toc_entry(self, text: str) -> bool:
        """
        判断是否为目录条目

        Args:
            text: 文本内容

        Returns:
            是否为目录条目
        """
        # 检查是否包含制表符后跟数字（页码）
        if '\t' in text and re.search(r'\t+\d+\s*$', text):
            return True

        # 检查是否是"标题 + 多个空格 + 页码"的模式
        if re.search(r'.+\s{3,}\d+\s*$', text):
            return True

        # 检查是否是纯数字结尾的短行（可能是页码）
        if re.match(r'^.{1,50}\s+\d+\s*$', text) and len(text.strip()) < 100:
            return True

        return False

    def _looks_like_toc_pattern(self, text: str, pattern: str) -> bool:
        """
        检查文本是否看起来像目录模式

        Args:
            text: 原始文本
            pattern: 匹配的模式

        Returns:
            是否看起来像目录条目
        """
        # 如果匹配的模式后面紧跟着制表符和数字，很可能是目录
        pattern_end = text.find(pattern) + len(pattern)
        if pattern_end < len(text):
            remaining = text[pattern_end:]
            if re.match(r'^\s*\t+\d+\s*$', remaining) or re.match(r'^\s{3,}\d+\s*$', remaining):
                return True

        return False

    def _extract_section_content(self, start_pos: int, end_pos: int,
                                 chapter: Dict, title_text: str) -> str:
        """
        提取章节内容

        Args:
            start_pos: 开始位置
            end_pos: 结束位置
            chapter: 章节信息
            title_text: 标题文本

        Returns:
            章节内容字符串
        """
        content_parts = []

        # 添加章节标题
        content_parts.append(title_text)

        # 提取内容
        for i in range(start_pos + 1, end_pos):
            if i >= len(self.paragraphs):
                break

            para_info = self.paragraphs[i]

            if para_info['type'] == 'paragraph':
                text = para_info['text']
                if text:
                    content_parts.append(text)

            elif para_info['type'] == 'table':
                table_content = self._format_table(para_info['element'])
                if table_content:
                    content_parts.append(table_content)

        return '\n\n'.join(content_parts)

    def _format_table(self, table: Table) -> str:
        """
        格式化表格内容

        Args:
            table: 表格对象

        Returns:
            格式化的表格字符串
        """
        if not table.rows:
            return ""

        content_parts = []
        content_parts.append("--- 表格内容 ---")

        # 提取表格数据
        table_data = []
        for row in table.rows:
            row_data = []
            for cell in row.cells:
                cell_text = cell.text.strip()
                row_data.append(cell_text)
            if any(row_data):  # 只添加非空行
                table_data.append(row_data)

        if not table_data:
            content_parts.append("--- 表格结束 ---")
            return '\n'.join(content_parts)

        # 计算列宽
        max_cols = max(len(row) for row in table_data)
        col_widths = [0] * max_cols

        for row in table_data:
            for i, cell in enumerate(row):
                if i < len(col_widths):
                    col_widths[i] = max(col_widths[i], len(cell))

        # 格式化表格行
        for row in table_data:
            formatted_row = []
            for i in range(max_cols):
                if i < len(row):
                    cell = row[i]
                else:
                    cell = ""
                formatted_row.append(cell)

            content_parts.append(" | ".join(formatted_row))

        content_parts.append("--- 表格结束 ---")

        return '\n'.join(content_parts)

    def get_all_chapters(self) -> List[Dict]:
        """
        获取所有章节信息

        Returns:
            章节列表
        """
        return self.chapters

    def get_chapter_tree(self) -> str:
        """
        获取章节树结构

        Returns:
            格式化的章节树
        """
        return self.outline_extractor.format_outline(self.chapters)


def main():
    """主函数，处理命令行参数"""
    parser = argparse.ArgumentParser(
        description="从 DOC/DOCX 文档中提取指定章节内容",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
使用示例:
  python from_server_docx_para.py 我的模板.docx --chapter 3.2
  python from_server_docx_para.py 报告.doc --chapter "第一章"
  python from_server_docx_para.py 文档.docx --tree

依赖要求:
  - 处理 .doc 文件需要安装 LibreOffice
  - 处理 .docx 文件仅需要 python-docx 库
        """
    )

    parser.add_argument(
        'file',
        help='要分析的 DOC/DOCX 文件路径'
    )

    parser.add_argument(
        '--chapter',
        type=str,
        help='要提取的章节编号，如 "3.2"'
    )

    parser.add_argument(
        '--tree',
        action='store_true',
        help='显示章节树结构'
    )

    args = parser.parse_args()

    try:
        # 1. 创建DocxParser实例
        doc_parser = DocxParser(args.file)

        if args.tree:
            # 显示章节树
            print("章节结构:")
            print("=" * 50)
            print(doc_parser.get_chapter_tree())

        elif args.chapter:
            # 2. 调用 get_chapter 方法来获取内容
            content = doc_parser.get_chapter(args.chapter)
            print(f"章节 {args.chapter} 内容:")
            print("=" * 50)
            print(content)

        else:
            print("请指定 --chapter 参数来提取章节内容，或使用 --tree 查看章节结构")

    except Exception as e:
        print(f"处理失败: {e}")
        import traceback
        traceback.print_exc()


if __name__ == "__main__":
    main()
