from gpu_server.Openai_Api_for_Qwen import *
import openai
openai.api_base = "http://116.62.63.204:8000/v1"

import copy

from docx import Document
llm = LLM_Qwen()

class LLM_Doc():
    def __init__(self, in_file_name):
        self.doc_name = in_file_name

        self.win32_doc = None
        self.win32_doc_app = None
        self.win32_constant = None
        
    def win32com_init(self):
        print(f'win32尝试打开文件"{self.doc_name}"')
        if is_win():
            import win32com.client as win32
            from win32com.client import constants
            self.win32_constant = constants
        else:
            print(f"未在windows系统下运行，无法获取word文档页码.")
            return

        import os

        self.win32_doc_app = win32.gencache.EnsureDispatch('Word.Application')  # 打开word应用程序
        self.win32_doc_app.Visible = False
        # doc_app.Visible = True
        curr_path = os.getcwd()
        self.win32_doc = self.win32_doc_app.Documents.Open(self.doc_name)
        print(f'win32打开了文件"{self.doc_name}"')

    def win32_close_file(self):
        if self.win32_doc:
            self.win32_doc_app.Documents.Close(self.doc_name)

    def get_paragraphs_generator_for_docx_file(self):
        try:
            doc = Document(self.doc_name)
        except Exception as e:
            print(f'文件"{self.doc_name}" 未找到。')
            return None

        for para in doc.paragraphs:
            yield para.text

    # 检查某doc文档所有para的错别词
    def check_wrong_written_in_docx_file(self):
        try:
            doc = Document(self.doc_name)
        except Exception as e:
            print(f'文件"{self.doc_name}" 未找到。')
            return

        p_result_list = []
        ii = 0
        for para in doc.paragraphs:
            ii += 1
            print(f"正在分析第{ii}个段落...")

            # para result登记
            print("*"*100)
            llm.ask_prepare(para.text).get_answer_and_sync_print()

# ============================关于角色提示============================
# 一、你希望llm了解哪些信息：
# 1) where are you based?
# 2) what do you do for work?
# 3) what are your hobbies and interests?
# 4) what subject can you talk about for hours?
# 5) what are some goals you have?
# 二、你希望llm怎样回复：
# 1) how formal or casual should llm be?
# 2) how long or short should responses generally be?
# 3) how do you want to be addressed?
# 4) should llm have opinions on topics or remain neutral?
# ============================关于角色提示============================
def main():
    llm = LLM_Qwen(history=True, history_max_turns=50, history_clear_method='pop')
    role_prompt = '你是一位错别字纠错专家。user发任何文本信息给你，你都需要找出文本中的所有错别字，并给出修改意见，修改意见要简单扼要；例如："我是你的好彭友"，你要指出"彭"字应为"朋"；如果没用错别字，你就返回"无错别字。" '
    llm.set_role_prompt(role_prompt)

    llm.ask_prepare('天气真好').get_answer_and_sync_print()

if __name__ == "__main__":
    main()