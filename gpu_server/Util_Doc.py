import sys
sys.path.append("..")
from singleton import AbstractSingleton
import pycorrector
from docx import Document

from dataclasses import dataclass, field
from typing import List

def is_win():
    import platform
    sys_platform = platform.platform().lower()
    if "windows" in sys_platform:
        return True
    else:
        return False

@dataclass(frozen = False)
class para_wrong_written_result:
    page_num: int = 0                                                   # para所在的页数
    paragraph_text: str = ''                                            # para的文本
    correct_suggestion_list: List[str] = field(default_factory=list)    # para中错别词的list

    def __post_init__(self):
        # __init__之后的处理
        pass

@dataclass(frozen = False)
class doc_wrong_written_result:
    doc_name: str = ''                                                      # 文件路径及名称
    result: List[para_wrong_written_result] = field(default_factory=list)   # 出错para的list

    def __post_init__(self):
        # __init__之后的处理
        pass

class Text_Topic_Search():
    def __init__(self, in_doc, in_search_keyword, in_context_window_length=30):
        self.doc = in_doc
        self.search_keyword = in_search_keyword
        self.context_window_length = in_context_window_length

        self._search_result_paragraphs = []

    def count(self):
        counting = False
        paras_num = 0
        # 开始统计，且统计数量<win_len
        for para in self.doc.paragraphs:
            # 本次统计完成，计数器置0
            if paras_num > self.context_window_length:
                paras_num = 0
                counting = False

            # 找到keyword（如标题中找到keyword"建设规模"，则先把标题append到结果中）
            if self.search_keyword in para.text:
                counting = True
                self._search_result_paragraphs.append(para.text)
                paras_num += 1
                continue

            if counting:
                self._search_result_paragraphs.append(para.text)
                paras_num += 1

        # 返回统计结果string
        return ' '.join(self._search_result_paragraphs)
        # return self._search_result_paragraphs


class Epdi_Text(AbstractSingleton):
    def __init__(self):
        # self.epdi_vocabulary = [
        #     '浙电经','电监会','赋码','导则','总平','总平图','国网',
        #     '孤网','网侧','主网','生技','主变','录波器','远动','遥信','遥调','主站','子站','表计','备调','端站',
        #     '已建','网损','综自','组屏','省调','地调','国调','污区','闪变','裕度',
        #     '对侧','开断','峰荷','腰荷','谷荷','幅值',
        #     '联变','排管','放坡','柔直',
        # ]

        self.corrector = pycorrector.correct
        self.doc_name = ''
        self.win32_doc = None
        self.win32_doc_app = None
        self.win32_constant = None

    def init(self, doc_name, init_corrector=False):
        self.doc_name = doc_name

        if init_corrector:
            # pycorrector参数的初始化
            pycorrector.enable_char_error(enable=False)     #关闭单个字的纠错
            pycorrector.set_custom_word_freq(path='./util_doc_custom_words.txt')

            # Win32com的初始化
            self.win32com_init()

    # 检查字符串的错别词
    def check_text(self, text):
        corrected_sentence, detail_list = self.corrector(text)
        return corrected_sentence, detail_list
        # detail_list格式为[(),(), ...], 其中()为(cur_item, corrected_item, begin_idx, end_idx)
        # 例如：[('鱼山', '玉山', 0, 2), ('二其', '二期', 11, 13), ('合记', '各级', 18, 20)]

    # 删除对电力设计院专有名词的误判
    # def del_check_result_in_voc(self, inout_text_errors_list):
    #     if len(inout_text_errors_list)==0:
    #         return
    #
    #     while True:
    #         index = 0
    #         has_epdi_voc = False
    #         for item in inout_text_errors_list:
    #             if item[0] in self.epdi_vocabulary:         # item[0]即为cur_item (cur_item, corrected_item, begin_idx, end_idx)
    #                 # print(f"voc {item} got")
    #                 del inout_text_errors_list[index]
    #                 has_epdi_voc = True
    #                 break
    #             index += 1
    #
    #         if has_epdi_voc==False:     # 如果没有词汇表重合情况，退出while
    #             break

    # win32com初始化
    def win32com_init(self):
        print(f'win32尝试打开文件"{self.doc_name}"')
        if is_win():
            import win32com.client as win32
            from win32com.client import constants
            self.win32_constant = constants
        else:
            print(f"未在windows系统下运行，无法获取word文档页码.")
            return

        import os

        self.win32_doc_app = win32.gencache.EnsureDispatch('Word.Application')  # 打开word应用程序
        self.win32_doc_app.Visible = False
        # doc_app.Visible = True
        curr_path = os.getcwd()
        self.win32_doc = self.win32_doc_app.Documents.Open(self.doc_name)
        print(f'win32打开了文件"{self.doc_name}"')

    # 获取para所在页码（仅限win32下，其他os只能将doc转换为pdf后获取para所在页码）
    def win32_get_para_page_num(self, text):
        if not is_win():
            return 0

        search_range = self.win32_doc.Content

        search_range.Find.Execute(FindText=text)
        search_range.Select()  # 这一步必须要有，否则获取的页码不对
        range = self.win32_doc_app.Selection.Range
        # print('绝对页：', rng.Information(constants.wdActiveEndPageNumber))
        # print('用户页：', rng.Information(constants.wdActiveEndAdjustedPageNumber))
        return range.Information(self.win32_constant.wdActiveEndAdjustedPageNumber)

    def win32_close_file(self):
        if self.win32_doc:
            self.win32_doc_app.Documents.Close(self.doc_name)

    def get_paragraphs_generator_for_docx_file(self):
        try:
            doc = Document(self.doc_name)
        except Exception as e:
            print(f'文件"{self.doc_name}" 未找到。')
            return None

        for para in doc.paragraphs:
            yield para.text

    # 检查某doc文档所有para的错别词
    def check_wrong_written_in_docx_file(self):
        try:
            doc = Document(self.doc_name)
        except Exception as e:
            print(f'文件"{self.doc_name}" 未找到。')
            return

        p_result_list = []
        ii = 0
        for para in doc.paragraphs:
            ii += 1
            corrected_sentence, text_errors_list = self.check_text(para.text)
            print(f"正在分析第{ii}个段落...")

            if len(text_errors_list)>0:
                # 从text_errors_list(cur_item, corrected_item, begin_idx, end_idx)抽取cur_item
                correct_suggestion_list = []
                for item in text_errors_list:
                    correct_suggestion = f'错别词"{item[0]}"建议改为"{item[1]}"，位于段落第{item[2]}个字符。'
                    correct_suggestion_list.append(correct_suggestion)

                # para result登记
                print("*"*100)
                p_page_num = self.win32_get_para_page_num(para.text)
                p_result = para_wrong_written_result(page_num=p_page_num, paragraph_text=para.text, correct_suggestion_list=correct_suggestion_list)
                print("段落建议：", p_result)
                p_result_list.append(p_result)

        # doc的所有para result统计
        rtn_result = doc_wrong_written_result(doc_name=self.doc_name, result=p_result_list)
        return rtn_result


if __name__ == '__main__':
    s_t = Epdi_Text()
    s_t.init("/Volumes/public/mbp15/mbp15_工作/===智慧能源====/200、===================科技项目===================/2023-08-07-LLM在能源电力系统咨询中的实战应用研究/LLM测试文档.docx")
    # s_t.init("Y:/mbp15/mbp15_prog/server/nps/gpu_server/LLM测试文档.docx")
    s_t.check_wrong_written_in_docx_file()
    s_t.win32_close_file()

    # s_t = Epdi_Text()
    # corrected_sent, detail = s_t.check_text('鱼山绿色石化基地一期、二其已经建成，合记设计用电负和达到1927MW。舟山绿色石化基地正在开展高性能树脂、高端新材料、炼化一体化改造提升项目相关前期工作，预计新增用电负荷约1000MW。')
    # print(corrected_sent, detail)
    # corrected_sent, detail = s_t.check_text('你是我的好彭友')
    # print(corrected_sent, detail)
    # corrected_sent, detail = s_t.check_text('你是我的好朋友')
    # print(corrected_sent, detail)


